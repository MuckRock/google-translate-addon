"""
DocumentCloud Add-On that translates documents using Google Translate services.
"""
import math
import os
import sys
from tempfile import NamedTemporaryFile

from documentcloud.addon import AddOn
from documentcloud.exceptions import APIError

# pylint: disable=no-name-in-module
from google.cloud import translate_v2 as translate

import docx


class Translate(AddOn):
    """DocumentCloud premium Add-On that translates documents"""

    def validate(self):
        """Validate that we can run the translation"""
        if self.get_document_count() is None:
            self.set_message(
                "It looks like no documents were selected. Search for some or "
                "select them and run again."
            )
            return False
        if not self.org_id:
            self.set_message("No organization to charge.")
            sys.exit(0)
        num_chars = 0
        for document in self.get_documents():
            num_chars += len(document.full_text)
        cost = math.ceil(num_chars / 75)
        try:
            self.charge_credits(cost)
        except ValueError:
            return False
        except APIError:
            return False
        return True

    def dry_run(self, documents):
        """ Calculates and displays cost to user, no real translation"""
        num_chars = 0
        for doc in documents:
            num_chars += len(doc.full_text)
        cost = math.ceil(num_chars / 75)
        self.set_message(
            f"There are {num_chars} characters in this document set. "
            f"It would cost {cost} AI credits to translate this document set."
        )
        sys.exit(0)

    def setup_credential_file(self):
        """Sets up Google Cloud credential file"""
        credentials = os.environ["TOKEN"]
        # put the contents into a named temp file
        # and set the var to the name of the file
        gac = NamedTemporaryFile(delete=False) # pylint: disable=consider-using-with
        gac.write(credentials.encode("ascii"))
        gac.close()
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = gac.name

    def translate_text(self, text, target_lan, source_lan):
        """Calls the translate client to do the translation"""
        translate_client = translate.Client()
        result = translate_client.translate(
            text, target_language=target_lan, source_language=source_lan
        )
        return result["translatedText"]

    def main(self):
        """ Grabs input and output languages, translates each document """
        if self.data.get("dry_run"):
            # If dry_run is selected, it will calculate the cost of translation.
            self.dry_run(self.get_documents())

        if not self.validate():
            self.set_message("You do not have sufficient AI credits to run this Add-On")
            sys.exit(0)

        # Sets up Google Cloud API Credential file
        self.setup_credential_file()
        # Retrieve input and out language charactor codes
        source_lang = self.data["input_lang"]
        target_lang = self.data["output_lang"]
        # Creates temporary directory out to store translations in before upload
        os.makedirs(os.path.dirname("./out/"), exist_ok=True)
        os.chdir("./out/")
        # For each document, translate the text and create a text file with the translation
        for document in self.get_documents():
            self.set_message(f"Translating {document.title}...")
            doc = docx.Document() # pylint: disable=no-member
            for page in range(1, document.page_count + 1):
                doc.add_heading(f"Page # {page} Translation", 3)
                translated_text = str(
                    self.translate_text(
                        document.get_page_text(page), target_lang, source_lang
                    )
                )
                doc.add_paragraph(translated_text)
                doc.add_page_break()
            doc.save(f"{document.title}-translation_{target_lang}.docx")
            self.set_message("Uploading translation...")
            # If project ID is specified, upload to that project.
            if self.data.get("project_id") is not None:
                self.client.documents.upload(
                    f"{document.title}-translation_{target_lang}.docx",
                    original_extension="docx",
                    title=f"{document.title}-translation_{target_lang}",
                    access=self.data["access_level"],
                    project=self.data.get("project_id"),
                )
            else:
                self.client.documents.upload(
                    f"{document.title}-translation_{target_lang}.docx",
                    original_extension="docx",
                    title=f"{document.title}-translation_{target_lang}",
                    access=self.data["access_level"],
                )


if __name__ == "__main__":
    Translate().main()
